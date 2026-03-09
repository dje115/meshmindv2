"""DOCX extraction using python-docx."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from ..models import (
    ExtractionMetadata,
    ExtractionStatus,
    NormalizedDocument,
    PageBlock,
)
from .base import ExtractionResult


def extract_docx(path: Path, provenance: dict) -> ExtractionResult:
    """Extract DOCX. Single logical page (no reliable page breaks in DOCX)."""
    try:
        doc = Document(str(path))
    except Exception as e:
        return ExtractionResult(failure_reason=f"docx read error: {e}")

    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    full_text = "\n\n".join(parts).strip()
    pages = [PageBlock(page_index=0, text=full_text)] if full_text else []
    meta = ExtractionMetadata(
        status=ExtractionStatus.SUCCESS,
        confidence=0.95,
        page_count=1 if full_text else 0,
        parser="python-docx",
    )
    return ExtractionResult(
        document=NormalizedDocument(
            full_text=full_text,
            pages=pages,
            extraction_metadata=meta,
            provenance=provenance,
        )
    )
